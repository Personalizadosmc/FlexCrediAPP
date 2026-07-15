from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

ROOT=Path(r"C:\AppMovil2026\FlexCrediAPP")
OUT=ROOT/"Memoria_Personal_Desarrollo_FlexCrediAPP.docx"
BLUE,ROYAL,GOLD,RED="0D2252","1A56DB","F59E0B","DC2626"
d=Document();s=d.sections[0]
s.top_margin=s.bottom_margin=Inches(.7);s.left_margin=s.right_margin=Inches(.78)
d.styles["Normal"].font.name="Aptos";d.styles["Normal"].font.size=Pt(11);d.styles["Normal"].paragraph_format.space_after=Pt(7);d.styles["Normal"].paragraph_format.line_spacing=1.15
for n,z,c in [("Title",30,BLUE),("Heading 1",20,BLUE),("Heading 2",15,ROYAL),("Heading 3",12,BLUE)]:
 st=d.styles[n];st.font.name="Aptos Display";st.font.size=Pt(z);st.font.bold=True;st.font.color.rgb=RGBColor.from_string(c)

def shade(c,color):
 pr=c._tc.get_or_add_tcPr();x=OxmlElement("w:shd");x.set(qn("w:fill"),color);pr.append(x)
def table(headers,rows,size=8.7):
 t=d.add_table(rows=1,cols=len(headers));t.style="Table Grid";t.alignment=WD_TABLE_ALIGNMENT.CENTER
 for i,h in enumerate(headers):
  c=t.rows[0].cells[i];c.text=h;shade(c,BLUE)
  for r in c.paragraphs[0].runs:r.font.bold=True;r.font.color.rgb=RGBColor(255,255,255);r.font.size=Pt(size)
 for j,row in enumerate(rows):
  cs=t.add_row().cells
  for i,v in enumerate(row):
   cs[i].text=str(v)
   if j%2:shade(cs[i],"F3F6FB")
   for p in cs[i].paragraphs:
    for r in p.runs:r.font.size=Pt(size)
 d.add_paragraph();return t
def bullets(xs):
 for x in xs:d.add_paragraph(x,style="List Bullet")
def numbers(xs):
 for x in xs:d.add_paragraph(x,style="List Number")
def note(code,text,kind="APP"):
 t=d.add_table(rows=1,cols=1);c=t.cell(0,0);shade(c,"FFF7E6");p=c.paragraphs[0]
 r=p.add_run(f"NOTA PARA CAPTURA {code} — {kind}\n");r.bold=True;r.font.color.rgb=RGBColor.from_string(GOLD);p.add_run(text)
 d.add_paragraph()
def code(text,title):
 p=d.add_paragraph();r=p.add_run(title);r.bold=True;r.font.color.rgb=RGBColor.from_string(ROYAL)
 t=d.add_table(rows=1,cols=1);c=t.cell(0,0);shade(c,"F1F5F9");r=c.paragraphs[0].add_run(text);r.font.name="Consolas";r.font.size=Pt(8)
 d.add_paragraph()
def toc():
 p=d.add_paragraph();r=p.add_run();b=OxmlElement("w:fldChar");b.set(qn("w:fldCharType"),"begin");i=OxmlElement("w:instrText");i.set(qn("xml:space"),"preserve");i.text='TOC \\o "1-3" \\h \\z \\u';e=OxmlElement("w:fldChar");e.set(qn("w:fldCharType"),"end");r._r.extend([b,i,e]);d.add_paragraph("En Word: clic derecho sobre el índice y seleccione Actualizar toda la tabla.")
def footer(p):
 p.alignment=WD_ALIGN_PARAGRAPH.RIGHT;r=p.add_run("FlexCrediAPP | ");r.font.size=Pt(8);b=OxmlElement("w:fldChar");b.set(qn("w:fldCharType"),"begin");i=OxmlElement("w:instrText");i.text="PAGE";e=OxmlElement("w:fldChar");e.set(qn("w:fldCharType"),"end");r._r.extend([b,i,e])
footer(s.footer.paragraphs[0]);s.header.paragraphs[0].text="MEMORIA PERSONAL DE DESARROLLO • FLEXCREDIAPP";s.header.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT

# Portada
d.add_paragraph();d.add_paragraph();logo=ROOT/"src"/"assets"/"flex-credi.png"
if logo.exists():
 p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run().add_picture(str(logo),width=Inches(1.5))
p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run("FLEXCREDIAPP");r.bold=True;r.font.size=Pt(34);r.font.color.rgb=RGBColor.from_string(BLUE)
p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run("Memoria personal del proceso de desarrollo");r.font.size=Pt(17);r.font.color.rgb=RGBColor.from_string(ROYAL)
p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run("Un reto individual de programación móvil híbrida");r.bold=True;r.font.color.rgb=RGBColor.from_string(GOLD)
table(["Dato","Información"],[("Institución","Universidad Abierta para Adultos (UAPA)"),("Asignatura","Programación de Dispositivos Móviles — ISW-307"),("Facilitador","Joan Manuel Gregorio Pérez"),("Participante","[ESCRIBIR NOMBRE COMPLETO]"),("Matrícula","[ESCRIBIR MATRÍCULA]"),("Fecha","[FECHA DE ENTREGA]"),("Proyecto","FlexCrediAPP 1.0.0 — Ionic + Angular + Capacitor")],9.5)
note("P-01","Colocar el logo oficial de la UAPA en la portada, conservando también el logo de FlexCredi.","PORTADA")
d.add_page_break();d.add_heading("Índice",1);toc();d.add_page_break()

d.add_heading("Presentación personal del proyecto",1)
d.add_paragraph("Decidí desarrollar FlexCrediAPP de manera individual porque quise asumir el proyecto como un reto personal y poner a prueba mis capacidades. Desde el principio entendí que no se trataba solamente de crear varias pantallas para cumplir una lista, sino de construir una aplicación con una idea clara, relacionar cada unidad de la asignatura con una necesidad real y comprobar que todo funcionara como una sola solución.")
d.add_paragraph("Elegí el tema de la gestión de préstamos porque permite aplicar de manera natural la navegación, el almacenamiento, los gestos, la conectividad, la ubicación, la cámara, el código QR, el audio, Bluetooth, NFC y los servicios web. Así nació FlexCrediAPP: una aplicación pensada para una persona o pequeño negocio que necesita registrar clientes, crear préstamos, calcular cuotas, recibir pagos, producir recibos, organizar tareas de cobro y mantener comunicación con sus clientes.")
d.add_paragraph("Trabajar individualmente significó asumir todos los roles: analizar los requisitos, diseñar la interfaz, programar, integrar plugins, corregir errores, probar en navegadores y preparar Android. Esto aumentó el nivel de dificultad, pero también me permitió comprender cómo cada parte depende de las demás. En esta memoria explico el proceso como lo viví, módulo por módulo, incluyendo mis decisiones, dificultades y aprendizajes.")
note("I-01","Captura del proyecto abierto en Visual Studio Code, mostrando el nombre FlexCrediAPP y las carpetas src, android e ios.","PROCESO")

d.add_heading("1. Problema, objetivos y alcance",1)
d.add_heading("1.1 Problema que quise resolver",2)
d.add_paragraph("Pensé en la realidad de un prestamista que anota clientes y pagos en libretas o conversa por WhatsApp sin tener un registro central. Cuando la información está separada es fácil olvidar una cuota, calcular mal un saldo o perder un comprobante. Por eso diseñé una app que concentrara el trabajo diario en el teléfono y que también pudiera aprovechar funciones propias del dispositivo.")
d.add_heading("1.2 Objetivo general",2)
d.add_paragraph("Mi objetivo fue desarrollar una aplicación móvil híbrida, funcional y organizada que permitiera administrar clientes, préstamos, cuotas, pagos, tareas y comunicaciones, cubriendo todas las unidades indicadas por el facilitador mediante Ionic, Angular y Capacitor.")
d.add_heading("1.3 Objetivos específicos",2)
bullets(["Crear una navegación clara con tabs y rutas protegidas.","Implementar CRUD persistente para los datos principales.","Calcular cuotas y estados de préstamos automáticamente.","Mostrar conectividad, GPS y clientes en un mapa.","Integrar audio, cámara, QR, BLE y NFC.","Consumir servicios web con carga, errores y alternativas locales.","Generar recibos y facilitar comunicación por WhatsApp y correo."])
d.add_heading("1.4 Alcance",2)
d.add_paragraph("La versión 1.0 incluye autenticación local, clientes, préstamos, pagos, tareas, noticias, mapa, podcasts, perfil, BLE, NFC, QR, recibos, PDF, correo, WhatsApp, historial y reportes. No pretende reemplazar todavía un sistema bancario ni un backend empresarial; los datos principales se conservan localmente para la demostración académica.")

d.add_heading("2. Cómo organicé el desarrollo",1)
d.add_paragraph("Antes de programar dividí el trabajo según las cinco agrupaciones propuestas por el facilitador. En vez de construir ejemplos separados, adapté cada grupo al tema financiero. También preparé servicios para que la lógica no quedara repetida en cada pantalla. Por ejemplo, DataService centraliza clientes, préstamos y pagos; NetworkService controla la red; GeolocationService obtiene la ubicación; y otros servicios se ocupan de cámara, QR, Bluetooth, NFC, correo y WhatsApp.")
d.add_paragraph("Utilicé componentes standalone y rutas con carga diferida para mantener organizado el proyecto. La navegación principal quedó formada por Inicio, Tareas, Novedades, Mapa y Perfil. Desde Inicio agregué accesos al flujo financiero: Nuevo préstamo, Recibir pago, Historial, Clientes y Reportes.")
table(["Etapa","Lo que realicé"],[("1. Análisis","Relacioné las diez unidades con un problema financiero real."),("2. Diseño","Definí colores, cards, tabs, formularios, estados y navegación."),("3. Base funcional","Construí autenticación, clientes, préstamos, cuotas y pagos."),("4. Módulos","Integré tareas, API, red, mapa, audio, cámara, QR, BLE y NFC."),("5. Servicios adicionales","Agregué correo, WhatsApp, recordatorios, PDF, historial y reportes."),("6. Pruebas","Ejecuté lint, build, pruebas unitarias y revisiones en distintos entornos.")])
note("O-01","Captura de app.routes.ts y del árbol pages/services para mostrar cómo organicé el proyecto.","CÓDIGO")

d.add_heading("3. Modelo de negocio de mi propuesta",1)
d.add_paragraph("Aunque el proyecto es académico, imaginé FlexCrediAPP como un producto para prestamistas y pequeños negocios. Su valor es reunir cobros, ubicación, comunicación y evidencias en una experiencia móvil sencilla.")
table(["Bloque Canvas","Mi propuesta"],[("Clientes","Prestamistas independientes, cobradores y microempresas."),("Valor","Control de cartera, menos errores y recibos inmediatos."),("Canales","APK Android, futura tienda y versión web."),("Relación","Autoservicio, ayuda, soporte y actualizaciones."),("Ingresos","Futuro plan gratuito y suscripción profesional."),("Recursos","Código, plugins, APIs, dispositivos y conocimiento."),("Actividades","Desarrollo, pruebas, mantenimiento y soporte."),("Socios","Ionic, Capacitor, OSM, EmailJS, WhatsApp y tiendas."),("Costos","Desarrollo, pruebas, publicación, nube y mantenimiento.")])
note("C-01","Insertar un Canvas visual con los nueve bloques anteriores.","DIAGRAMA")

d.add_heading("4. Diseño y experiencia visual",1)
d.add_paragraph("Quise que la aplicación se sintiera financiera y profesional. Elegí azul oscuro como color principal porque comunica confianza, azul brillante para acciones, dorado como acento, verde para pagos correctos y rojo para atrasos o eliminación. Usé cards con bordes redondeados, iconos de Ionicons, sombras suaves, botones amplios y mensajes breves.")
d.add_paragraph("También presté atención a la retroalimentación. Cuando la app está cargando muestro un spinner; cuando una acción termina presento un toast; cuando se va a eliminar o registrar un pago utilizo una confirmación; y cuando no existen datos muestro un estado vacío con una acción clara. Estas decisiones fueron importantes porque una app móvil debe comunicar siempre qué está ocurriendo.")
table(["Color","Uso"],[("#0D2252","Identidad y encabezados."),("#1A56DB","Botones y navegación activa."),("#F59E0B","Acentos financieros."),("#059669","Éxito y pagos."),("#DC2626","Atrasos y acciones peligrosas."),("#EEF2FF","Fondos suaves.")])
note("D-01","Captura del Dashboard donde se aprecien colores, cards, iconos y tabs.","APP")
note("D-02","Captura de uno o varios wireframes iniciales; si no existen, recrear bocetos simples de Login, Inicio, Tareas, Mapa y Perfil.","DISEÑO")

d.add_page_break();d.add_heading("5. Desarrollo paso a paso por módulos",1)
d.add_paragraph("A continuación describo cómo desarrollé cada módulo y cómo adapté los requisitos a FlexCrediAPP.")

d.add_heading("5.1 Módulo 1 — Navegación, acceso y novedades",1)
d.add_heading("Cómo comencé este módulo",2)
d.add_paragraph("Comencé creando la estructura que sostendría toda la aplicación. Primero preparé las pantallas de registro e inicio de sesión. El usuario escribe su información, la app valida los campos y, cuando los datos son correctos, guarda la sesión local. Después agregué AuthGuard para impedir que una persona sin sesión pudiera entrar directamente a las rutas internas.")
note("M1-01","Captura de Login o Registro funcionando. Debajo colocar una captura corta de login.page.html o login.page.ts con el formulario y la validación.","APP + CÓDIGO")
d.add_heading("Construcción de la navegación",2)
d.add_paragraph("Luego construí la navegación con ion-tabs. Seleccioné cinco secciones principales: Inicio, Tareas, Novedades, Mapa y Perfil. Cada tab conduce a una página independiente y el router administra las rutas. En el Dashboard añadí accesos a las operaciones financieras que no debían ocupar una pestaña fija: clientes, nuevo préstamo, recibir pago, historial y reportes.")
d.add_paragraph("Durante las pruebas descubrí que Reportes retenía una transición y requería dos toques. Este problema también bloqueaba temporalmente los accesos inferiores. Lo investigué como un problema de navegación entre el outlet de los tabs y una pantalla independiente. Finalmente utilicé navegación de Ionic sin animación para esa ruta y la dejé cargada directamente, evitando que el gesto quedara retenido.")
note("M1-02","Captura del Dashboard completo y otra de tabs.page.html o app.routes.ts mostrando las cinco pestañas y las rutas.","APP + CÓDIGO")
d.add_heading("Consumo de noticias con API REST",2)
d.add_paragraph("Para cubrir servicios web desarrollé Novedades. Es importante aclarar que Ionic no proporciona una API de noticias; utilicé HttpClient de Angular dentro de la aplicación Ionic para consumir JSONPlaceholder, una API pública utilizada para simular publicaciones. Transformé los posts para presentarlos como novedades y consejos financieros, agregando categorías, iconos y fechas.")
d.add_paragraph("El proceso fue: mostrar el estado de carga, realizar GET, limitar los resultados, transformar cada registro y actualizar las cards. También configuré un timeout de ocho segundos y catchError. Si la API no responde, la app muestra noticias locales de respaldo. Añadí pull-to-refresh y un botón de reintento para que el módulo siguiera siendo útil sin conexión.")
code("GET https://jsonplaceholder.typicode.com/posts\nHttpClient → timeout → map → catchError → noticias locales","Flujo simplificado de la API REST")
note("M1-03","Captura de Novedades con cards y pull-to-refresh. En código mostrar NoticiasService.obtenerNoticias con HttpClient, timeout, map y catchError.","APP + CÓDIGO")
d.add_heading("Lo que aprendí",2)
d.add_paragraph("Este módulo me enseñó que una navegación aparentemente sencilla puede tener detalles táctiles importantes y que consumir una API no es solamente mostrar datos: también hay que manejar espera, pérdida de Internet, errores y una alternativa para el usuario.")

d.add_page_break();d.add_heading("5.2 Módulo 2 — Tareas, interfaces, gestos y almacenamiento",1)
d.add_heading("Planteamiento del módulo",2)
d.add_paragraph("Adapté la lista de tareas al trabajo de cobranza. Una tarea puede ser una llamada, una visita, un cobro o una actividad general. Le agregué título, descripción, cliente opcional, prioridad, fecha límite y estado. El objetivo fue que el usuario pudiera preparar su jornada sin depender de otra aplicación.")
d.add_heading("CRUD y persistencia",2)
d.add_paragraph("Creé TareasService para separar el almacenamiento de la pantalla. Implementé las cuatro operaciones CRUD: crear una tarea, leer la lista, actualizar y eliminar. Elegí Ionic Storage porque ofrece una API asíncrona y en navegador trabaja sobre almacenamiento persistente como IndexedDB. Cada vez que una tarea cambia, guardo la colección completa con su orden actualizado.")
note("M2-01","Captura del modal Nueva tarea y captura de TareasService mostrando agregarTarea, actualizarTarea y eliminarTarea.","APP + CÓDIGO")
d.add_heading("Gestos e interacción",2)
d.add_paragraph("Después añadí los gestos solicitados. Con ion-item-sliding el usuario desliza una tarea para editar, cancelar o eliminar. Con ion-refresher puede recargar la lista. Con ion-reorder-group cambia el orden de los pendientes. También programé una pulsación prolongada: al mantener presionada una tarea se abre la edición. Debí controlar la propagación de eventos para que pulsar Completar no activara también la tarjeta.")
d.add_paragraph("Dividí la vista en Pendientes, Completadas y Canceladas. Agregué estadísticas y colores de prioridad para que el usuario pueda entender el estado sin leer todos los detalles.")
note("M2-02","Captura de una tarea deslizada mostrando las opciones. En código mostrar ion-item-sliding e ion-item-options.","APP + CÓDIGO")
note("M2-03","Captura de reorder o pull-to-refresh y fragmento de onReorder/onPressStart.","APP + CÓDIGO")
d.add_heading("Dificultad personal",2)
d.add_paragraph("La dificultad principal fue combinar varios gestos en el mismo elemento sin que se activaran dos acciones. Resolví esto usando temporizadores cancelables, touchstart/touchend y stopPropagation. Comprendí que en móvil no basta con que un botón funcione con el mouse; también debe responder correctamente al dedo.")

d.add_page_break();d.add_heading("5.3 Módulo 3 — Conectividad, ubicación y mapa",1)
d.add_heading("Estado de la conexión",2)
d.add_paragraph("Primero implementé NetworkService con @capacitor/network. Obtengo el estado inicial y escucho networkStatusChange. Publiqué el resultado mediante BehaviorSubject para que cualquier pantalla pueda actualizarse. En el Dashboard muestro si hay conexión, el tipo de red y una cola de cambios pendientes. En navegador añadí eventos online/offline como respaldo.")
note("M3-01","Captura del indicador de red online y otra offline. En código mostrar BehaviorSubject y networkStatusChange.","APP + CÓDIGO")
d.add_heading("Ubicación actual",2)
d.add_paragraph("Después trabajé con @capacitor/geolocation. La app solicita permiso y busca una posición de alta precisión con un timeout de veinte segundos. Si se ejecuta en navegador utiliza navigator.geolocation. Este fue uno de los puntos más variables: en algunos navegadores o dispositivos la ubicación fallaba por permisos, precisión desactivada o falta de HTTPS, mientras que en otros funcionaba de forma excelente.")
d.add_heading("Mapa y clientes",2)
d.add_paragraph("Construí el mapa con Leaflet y mosaicos de OpenStreetMap. Inicialmente se centra en República Dominicana. Cada cliente puede tener latitud y longitud; si solo tiene dirección, consulto Nominatim para geocodificarla. Después creo un marcador con las iniciales y color del cliente. Al tocarlo aparece un popup con teléfono, dirección y cantidad de préstamos activos.")
d.add_paragraph("También calculo distancia con la fórmula Haversine para contar clientes dentro de tres kilómetros. Esto convierte el requisito de geolocalización en una función útil para planificar cobros.")
note("M3-02","Captura del mapa con ubicación actual y varios marcadores. En código mostrar L.map, tileLayer y getCurrentPosition.","APP + CÓDIGO")
note("M3-03","Captura de un popup de cliente y fragmento donde se crea el marcador personalizado.","APP + CÓDIGO")
d.add_heading("Cómo enfrenté los fallos",2)
d.add_paragraph("Añadí mensajes diferentes para permiso denegado, navegador sin contexto seguro y falta de ubicación precisa. También esperé a que el contenedor estuviera listo antes de crear Leaflet y actualicé los marcadores al volver a la pantalla. Esta experiencia me confirmó que la geolocalización siempre debe probarse en varios entornos.")

d.add_page_break();d.add_heading("5.4 Módulo 4 — Podcasts, cámara y código QR",1)
d.add_heading("Podcasts simulados",2)
d.add_paragraph("Para multimedia creé una sección de podcasts financieros. Ionic tampoco ofrece una API de podcasts; utilicé la HTML5 Audio API dentro de la app Ionic y direcciones públicas de SoundHelix para simular episodios. Las portadas se cargan desde Picsum. El contenido es demostrativo y permite comprobar la reproducción remota.")
d.add_paragraph("Programé play, pausa, stop, anterior, siguiente, barra de progreso y volumen. Los eventos timeupdate, canplay, ended, play y pause mantienen la interfaz sincronizada con el elemento audio. Elegí ion-range para que el usuario pueda mover la reproducción y cambiar volumen.")
note("M4-01","Captura del catálogo y reproductor. En código mostrar el elemento audio y los listeners timeupdate/ended.","APP + CÓDIGO")
d.add_heading("Cámara y galería",2)
d.add_paragraph("Para la cámara creé CameraService con @capacitor/camera. La foto se solicita como Data URL para guardarla en el perfil. En web intenté getUserMedia con vista previa y, si no está disponible, utilizo un input de archivo. Esto permite cambiar foto personal y logo de empresa en distintos entornos.")
note("M4-02","Captura del perfil cambiando foto o logo. En código mostrar Camera.getPhoto y el fallback web.","APP + CÓDIGO")
d.add_heading("Escaneo QR aplicado al cobro",2)
d.add_paragraph("Integré el QR en Recibir pago. En Android uso ML Kit Barcode Scanning y en navegador intento BarcodeDetector con la cámara. Cuando se lee un valor, busco coincidencia por cliente o préstamo y selecciono automáticamente la cuota. Al cerrar el escáner detengo las pistas de video para no dejar la cámara encendida.")
note("M4-03","Captura del marco QR y, si es posible, del cliente cargado. En código mostrar QrService y procesarQr.","APP + CÓDIGO")
d.add_heading("Dificultades",2)
d.add_paragraph("El audio remoto necesita Internet y la reproducción debe comenzar por una acción del usuario. Cámara y QR cambian mucho entre web y Android. Por eso mantuve rutas nativas y alternativas web, controlé permisos, cancelación y liberación de recursos.")

d.add_page_break();d.add_heading("5.5 Módulo 5 — Bluetooth, NFC, perfil y configuración",1)
d.add_heading("Construcción del perfil",2)
d.add_paragraph("Usé Perfil como centro de identidad y configuración. El usuario puede cambiar nombre, empresa, datos personales, foto y logo. También puede activar tema oscuro, notificaciones, sonidos, sincronización y vibración. Guardé estas preferencias con Ionic Storage para que permanezcan después de cerrar la app.")
note("M5-01","Captura completa del Perfil y fragmento de ProfileService con getConfig/saveConfig.","APP + CÓDIGO")
d.add_heading("Bluetooth Low Energy",2)
d.add_paragraph("Creé BluetoothService con @capacitor-community/bluetooth-le. Antes de escanear detecto si la plataforma es nativa, web o no soportada. Los resultados muestran nombre, identificador y RSSI; traduzco la señal a excelente, buena, débil o muy débil. Después permito intentar una conexión.")
note("M5-02","Captura del modal BLE con dispositivos y RSSI. En código mostrar escanear y conectar.","APP + CÓDIGO")
d.add_heading("NFC",2)
d.add_paragraph("Con NfcService verifico si el hardware está disponible, puedo escribir texto o información del recibo y también leer contenido. Registro cada acción en un historial indicando tipo, destino, fecha y resultado. Como NFC y BLE dependen del hardware, evito presentar una simulación como si fuera una operación real.")
note("M5-03","Captura de lectura/escritura NFC y del historial. En código mostrar verificarSoporteNativo y escribir/leer.","APP + CÓDIGO")
d.add_heading("Aprendizaje",2)
d.add_paragraph("Este módulo me enseñó la diferencia entre una función que se puede demostrar en navegador y otra que necesita obligatoriamente un teléfono compatible. La validación final de BLE y NFC debe hacerse en Android físico con permisos y sensores activados.")

d.add_page_break();d.add_heading("6. Funciones financieras que añadí",1)
d.add_heading("6.1 Clientes y préstamos",2)
d.add_paragraph("Después de construir la base, desarrollé el flujo real de FlexCredi. En Clientes implementé registro, búsqueda, edición, eliminación, contacto, dirección e historial. En Nuevo préstamo selecciono un cliente y utilizo un formulario reactivo para validar monto, interés, número de cuotas y fecha.")
d.add_paragraph("DataService calcula el total y divide las cuotas. Según la frecuencia suma 7, 15 o 30 días. Cada cuota inicia pendiente y el préstamo puede estar activo, atrasado o completado. Esta relación fue esencial para que Dashboard, pagos e informes compartieran la misma información.")
note("F-01","Captura de Nuevo cliente y Nuevo préstamo. En código mostrar calcularCuotas.","APP + CÓDIGO")
d.add_heading("6.2 Recibir pago y recibo",2)
d.add_paragraph("Recibir pago carga préstamos activos, permite buscar o usar QR y selecciona la primera cuota pendiente o atrasada. Antes de guardar muestro una confirmación. Al aceptar, registro el pago, marco la cuota como pagada y, si no quedan cuotas, completo el préstamo.")
d.add_paragraph("Después muestro un recibo profesional con empresa, logo, cliente, número, fecha, monto, cuota y saldo. También programé la generación de PDF, descarga, Web Share, WhatsApp y NFC. Este fue uno de los flujos más grandes porque conecta datos, comunicación y evidencia.")
note("F-02","Captura de selección de cuota y del recibo final. En código mostrar registrarPago/procesar y generarPdfFactura.","APP + CÓDIGO")
d.add_heading("6.3 WhatsApp, correo y recordatorios",2)
d.add_paragraph("Añadí WhatsappService para normalizar teléfonos dominicanos y construir mensajes de recordatorio o recibo. La aplicación abre wa.me con el texto preparado; el usuario conserva control y pulsa Enviar en WhatsApp.")
d.add_paragraph("Para correo integré EmailJS. Al registrar un pago preparo los parámetros de la plantilla: cliente, cuota, monto, empresa, saldo y estado. Cargo el SDK, envío el mensaje y muestro si fue correcto. EmailJS es un servicio externo adicional, no un requisito mínimo, pero aporta valor al proyecto.")
d.add_paragraph("ReminderService revisa fechas y prepara recordatorios previos, del día o atrasados. El Dashboard muestra el conteo y permite abrir el flujo de cobro.")
note("F-03","Captura de WhatsApp con mensaje prellenado y correo recibido. Ocultar datos reales. En código mostrar WhatsappService y EmailService.","APP + CÓDIGO")
d.add_heading("6.4 Historial y reportes",2)
d.add_paragraph("El Historial permite ver pagos globales o agrupados por cliente. Reportes calcula cartera, clientes, atrasos, préstamos completados y cobrado en semana, mes o año. También preparé una gráfica semanal. Estas pantallas muestran que los registros no quedan aislados, sino que se convierten en información para tomar decisiones.")
note("F-04","Captura de Historial y Reportes con indicadores y gráfica.","APP")

d.add_heading("7. Almacenamiento y flujo de información",1)
d.add_paragraph("Utilicé localStorage para usuarios, clientes, préstamos, cuotas, pagos y cola offline. Para tareas, perfil, fotos, logo, configuraciones e historial de compartidos utilicé Ionic Storage. Esta combinación fue suficiente para una versión académica sin servidor y me permitió demostrar persistencia.")
table(["Dato","Cómo lo guardé"],[("Usuarios/sesión","localStorage"),("Clientes","fc_clientes"),("Préstamos/cuotas","fc_prestamos"),("Pagos","fc_pagos"),("Cola offline","fc_sync_queue"),("Tareas/perfil/configuración","Ionic Storage")])
d.add_paragraph("El flujo más representativo es el pago: selecciono préstamo, registro Pago, actualizo Cuota, recalculo Préstamo y luego Historial, Dashboard y Reportes leen el nuevo estado. Esta relación me obligó a cuidar que una sola acción no produjera datos inconsistentes.")
note("AL-01","Captura de models/index.ts y DataService mostrando Cliente, Prestamo, Cuota y Pago.","CÓDIGO")

d.add_heading("8. Pruebas, errores y proceso de depuración",1)
d.add_paragraph("Durante el desarrollo no me limité a comprobar el camino ideal. Probé navegación, formularios vacíos, ausencia de datos, pérdida de Internet, permisos denegados y diferentes tamaños de pantalla. Ejecuté npm run lint, npm run build y pruebas Karma. También sincronicé Android con Capacitor.")
table(["Prueba","Resultado"],[("Registro/login y rutas protegidas","Correcto."),("CRUD de clientes/tareas","Persistencia correcta."),("Cálculo de préstamo/cuotas","Correcto."),("Registro de pago/recibo","Correcto."),("API noticias online/offline","API o respaldo local."),("Mapa y marcadores","Correcto con red."),("GPS","Variable por permisos/navegador."),("Podcasts","Correcto con red."),("QR/cámara","Validación web y Android."),("BLE/NFC","Requiere Android físico."),("Lint/build/test","Finalizados correctamente.")])
d.add_heading("Errores que tuve que resolver",2)
bullets(["Mensajes con caracteres dañados: normalicé codificación UTF-8.","Spinners que no desaparecían: cerré estados en éxito y error.","Recibir pago quedó visualmente vacío durante una limpieza: restauré el bloque sin perder datos.","Reportes necesitaba dos toques: corregí la transición entre tabs y página independiente.","GPS falló en algunos entornos: agregué timeout, fallback y mensajes.","Gestos de tareas chocaban: controlé temporizadores y propagación."])
note("PR-01","Captura de terminal con lint, build y TOTAL SUCCESS.","PRUEBAS")
note("PR-02","Captura de la app en navegador, emulador Android y teléfono real.","PRUEBAS")

d.add_heading("9. Dificultades y reflexión personal",1)
d.add_paragraph("La dificultad más grande fue comprobar que la aplicación respondiera de forma parecida en diferentes navegadores y dispositivos. En una computadora una función podía trabajar bien y en un teléfono pedir permisos distintos. La ubicación fue el mejor ejemplo: en ciertos casos falló y en otros mostró una precisión excelente. Lo mismo ocurre con cámara, QR, BLE y NFC.")
d.add_paragraph("Crear una aplicación híbrida es un reto grande. Aunque se comparte gran parte del código, cada plataforma conserva reglas propias. Por eso estas aplicaciones toman tiempo: hay que revisar navegación, permisos, sensores, conectividad, almacenamiento, orientación y rendimiento. Aprendí que compilar no significa terminar; una aplicación se considera preparada cuando se prueba repetidamente y se corrigen los detalles que afectan al usuario.")
d.add_paragraph("Asumir el proyecto individualmente fue exigente porque tuve que cambiar constantemente de perspectiva: usuario, diseñador, programador y probador. Sin embargo, fue una decisión valiosa porque pude comprender el proyecto completo y no solamente una sección.")

d.add_heading("10. Conclusiones y mejoras futuras",1)
d.add_paragraph("Considero que FlexCrediAPP cumple el propósito del proyecto porque integra las diez unidades en un tema único. No construí cinco ejemplos desconectados; cada módulo participa en el proceso de administrar y cobrar un préstamo. Navegación, storage, red, GPS, audio, cámara, QR, BLE, NFC y API REST tienen una función identificable dentro de la app.")
d.add_paragraph("Para una versión futura incorporaría un backend seguro, múltiples usuarios por empresa, sincronización real, base cifrada, recuperación de contraseña, notificaciones push, reportes exportables, copias de seguridad y publicación firmada en Google Play. También ampliaría las pruebas automáticas y la matriz de dispositivos.")

d.add_heading("11. Referencias",1)
refs=["Angular. (2026). Angular documentation. https://angular.dev/","Capacitor. (2026). Capacitor documentation. https://capacitorjs.com/docs","EmailJS. (2026). EmailJS documentation. https://www.emailjs.com/docs/","Google. (2026). ML Kit barcode scanning. https://developers.google.com/ml-kit/vision/barcode-scanning","Ionic. (2026). Ionic Framework documentation. https://ionicframework.com/docs","Leaflet. (2025). Leaflet documentation. https://leafletjs.com/","MDN Web Docs. (2025). HTMLMediaElement. https://developer.mozilla.org/docs/Web/API/HTMLMediaElement","MDN Web Docs. (2025). MediaDevices.getUserMedia(). https://developer.mozilla.org/docs/Web/API/MediaDevices/getUserMedia","OpenStreetMap contributors. (2026). OpenStreetMap. https://www.openstreetmap.org/","OpenStreetMap Foundation. (2026). Nominatim documentation. https://nominatim.org/","ReactiveX. (2026). RxJS. https://rxjs.dev/","TypeScript. (2026). TypeScript documentation. https://www.typescriptlang.org/docs/"]
for ref in refs:
 p=d.add_paragraph(ref);p.paragraph_format.left_indent=Inches(.3);p.paragraph_format.first_line_indent=Inches(-.3)

d.add_heading("12. Anexos",1)
d.add_heading("Anexo A — Guion personal para la demostración",2)
numbers(["Explico por qué elegí el tema y trabajar individualmente.","Inicio sesión y muestro tabs.","Registro cliente y préstamo.","Demuestro tareas y gestos.","Muestro red, noticias, GPS y mapa.","Reproduzco un podcast.","Pruebo cámara/QR y, en teléfono compatible, BLE/NFC.","Registro pago y comparto recibo.","Cierro con historial, reportes y aprendizaje."])
d.add_heading("Anexo B — Lista resumida de capturas",2)
table(["Sección","Capturas"],[("Inicio","Proyecto, Login, Dashboard y rutas."),("Módulo 1","Novedades y código HttpClient."),("Módulo 2","Tareas, swipe, reorder y storage."),("Módulo 3","Red, GPS, mapa y marcadores."),("Módulo 4","Podcasts, cámara y QR."),("Módulo 5","Perfil, BLE y NFC."),("Finanzas","Cliente, préstamo, pago, recibo, correo, WhatsApp, historial y reportes."),("Pruebas","Lint, build, test, navegador, emulador y teléfono.")])
d.add_heading("Anexo C — Instalación",2)
numbers(["Ejecutar npm install.","Navegador: npm start o ionic serve.","Construir: npm run build.","Sincronizar: npx cap sync.","Abrir Android: npx cap open android.","Generar APK en Android Studio con Build > Build APK(s)."])
note("FINAL","Antes de entregar: completar datos personales, insertar capturas, actualizar índice, numerar figuras, revisar ortografía y ocultar información sensible.","REVISIÓN")

d.core_properties.title="Memoria personal de desarrollo de FlexCrediAPP";d.core_properties.author="[NOMBRE DEL PARTICIPANTE]";d.core_properties.subject="Proyecto individual ISW-307"
d.save(OUT);print(OUT)

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

ROOT = Path(r"C:\AppMovil2026\FlexCrediAPP")
OUT = ROOT / "Documentacion_Tecnica_FlexCrediAPP.docx"
BLUE, ROYAL, GOLD, GREEN, RED, GRAY = "0D2252", "1A56DB", "F59E0B", "059669", "DC2626", "64748B"

d = Document()
s = d.sections[0]
s.top_margin, s.bottom_margin = Inches(.65), Inches(.65)
s.left_margin, s.right_margin = Inches(.72), Inches(.72)
d.styles["Normal"].font.name = "Aptos"
d.styles["Normal"].font.size = Pt(10.5)
d.styles["Normal"].paragraph_format.space_after = Pt(6)
for n, z, c in [("Title", 30, BLUE), ("Heading 1", 20, BLUE), ("Heading 2", 15, ROYAL), ("Heading 3", 12, BLUE)]:
    st = d.styles[n]
    st.font.name, st.font.size, st.font.bold = "Aptos Display", Pt(z), True
    st.font.color.rgb = RGBColor.from_string(c)

def shade(cell, color):
    pr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd"); sh.set(qn("w:fill"), color); pr.append(sh)

def table(headers, rows, size=8.3):
    t = d.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c,BLUE)
        for r in c.paragraphs[0].runs: r.font.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(size)
    for j,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
            if j%2: shade(cells[i],"F4F7FC")
            for p in cells[i].paragraphs:
                for r in p.runs:r.font.size=Pt(size)
    d.add_paragraph()
    return t

def bullets(items):
    for x in items: d.add_paragraph(x, style="List Bullet")

def numbers(items):
    for x in items: d.add_paragraph(x, style="List Number")

def note(title, text, red=False):
    t=d.add_table(rows=1,cols=1); c=t.cell(0,0); shade(c,"FDECEC" if red else "FFF7E6")
    p=c.paragraphs[0]; r=p.add_run(title+"\n"); r.bold=True; r.font.color.rgb=RGBColor.from_string(RED if red else GOLD); p.add_run(text)
    d.add_paragraph()

def capture(code, text, kind="INTERFAZ"):
    note(f"CAPTURA PENDIENTE {code} — {kind}", text+" Incluya número de figura, título, dispositivo/navegador y una explicación breve. No deje este recuadro vacío en la entrega final.")

def code(text, title):
    p=d.add_paragraph(); r=p.add_run(title); r.bold=True; r.font.color.rgb=RGBColor.from_string(ROYAL)
    t=d.add_table(rows=1,cols=1); c=t.cell(0,0); shade(c,"F1F5F9")
    p=c.paragraphs[0]; rr=p.add_run(text); rr.font.name="Consolas"; rr.font.size=Pt(8); rr.font.color.rgb=RGBColor.from_string(BLUE)
    d.add_paragraph()

def toc():
    p=d.add_paragraph(); r=p.add_run(); b=OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"),"begin")
    i=OxmlElement("w:instrText"); i.set(qn("xml:space"),"preserve"); i.text='TOC \\o "1-3" \\h \\z \\u'
    sep=OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"),"separate"); tx=OxmlElement("w:t"); tx.text="Clic derecho aquí > Actualizar campo > Actualizar toda la tabla."; sep.append(tx)
    e=OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"),"end"); r._r.extend([b,i,sep,e])

def footer_page(p):
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=p.add_run("FlexCrediAPP  |  "); r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRAY)
    b=OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"),"begin"); i=OxmlElement("w:instrText"); i.text="PAGE"; e=OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"),"end"); r._r.extend([b,i,e])

s.header.paragraphs[0].text="FLEXCREDIAPP  •  DOCUMENTACIÓN TÉCNICA Y FUNCIONAL"
s.header.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in s.header.paragraphs[0].runs:r.font.size=Pt(8);r.font.bold=True;r.font.color.rgb=RGBColor.from_string(ROYAL)
footer_page(s.footer.paragraphs[0])

# Portada
d.add_paragraph(); d.add_paragraph()
logo=ROOT/"src"/"assets"/"flex-credi.png"
if logo.exists():
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run().add_picture(str(logo),width=Inches(1.5))
p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run("FLEXCREDIAPP");r.bold=True;r.font.size=Pt(34);r.font.color.rgb=RGBColor.from_string(BLUE)
p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run("Sistema móvil para gestión de préstamos, clientes y cobros");r.font.size=Pt(16);r.font.color.rgb=RGBColor.from_string(ROYAL)
p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run("DOCUMENTACIÓN TÉCNICA, FUNCIONAL Y DE PRUEBAS");r.bold=True;r.font.color.rgb=RGBColor.from_string(GOLD)
table(["Dato","Información"],[
    ("Institución","Universidad Abierta para Adultos (UAPA)"),("Asignatura","Programación de Dispositivos Móviles — ISW-307"),
    ("Facilitador","Joan Manuel Gregorio Pérez"),("Integrante(s)","[COMPLETAR NOMBRE(S) Y MATRÍCULA(S)]"),
    ("Carrera / recinto","[COMPLETAR]"),("Fecha de entrega","[COMPLETAR FECHA OFICIAL]"),("Versión","FlexCrediAPP 1.0.0")],9.5)
capture("P-01","Inserte el logo oficial UAPA en la portada junto al logo de FlexCredi.","PORTADA")
d.add_page_break()

d.add_heading("Control y preparación del documento",1)
table(["Campo","Valor"],[("Título","Documentación técnica y funcional de FlexCrediAPP"),("Estado","Final para revisión académica"),("Stack","Ionic 8, Angular 20, Capacitor 8, TypeScript 5.9"),("Plataformas","Web, Android e iOS"),("Repositorio","[COMPLETAR URL, SI APLICA]")])
d.add_heading("Datos que deben completarse",2)
bullets(["Nombres, matrículas, carrera, recinto y fecha oficial.","Responsable real de cada módulo.","Logo oficial UAPA y todas las capturas indicadas.","Índice, números de página, figuras y tablas actualizados en Word."])
note("IMPORTANTE","El contenido ya está adaptado al código real. Sustituya campos entre corchetes y no entregue recuadros de capturas vacíos.")
d.add_heading("Índice",1);toc();d.add_page_break()

d.add_heading("Resumen ejecutivo",1)
d.add_paragraph("FlexCrediAPP es una aplicación híbrida para prestamistas independientes y pequeños negocios de crédito. Centraliza clientes, préstamos, cuotas, cobros, recibos, tareas y reportes, y aprovecha capacidades móviles como GPS, cámara, QR, Bluetooth y NFC. La solución cubre las diez unidades del curso mediante una experiencia coherente orientada a la cobranza.")
d.add_paragraph("Además del mínimo académico, incorpora autenticación local, cálculos automáticos, recordatorios, WhatsApp, correo EmailJS, recibos PDF, historial, geocodificación con Nominatim, OpenStreetMap, operación local y mecanismos alternativos para navegador.")
table(["Módulo","Unidades","Implementación"],[("1. Navegación y novedades","1, 2, 10","Tabs, AuthGuard, routing, API REST y respaldo local."),("2. Tareas e interfaces","3, 9","CRUD, swipe, long press, reorder, refresher e Ionic Storage."),("3. Conectividad y mapa","4, 6","Network, GPS, Leaflet, OSM, Nominatim y marcadores."),("4. Multimedia/cámara/QR","7, 8","Audio personalizado, cámara, galería y QR híbrido."),("5. BLE/NFC/perfil","5, 9","BLE, NFC, configuración, fotos e historial.")])

d.add_heading("1. Introducción y objetivos",1)
d.add_heading("1.1 Problema",2)
d.add_paragraph("La información de pequeños prestamistas suele quedar dispersa entre libretas, hojas de cálculo, chats y comprobantes. Esto dificulta conocer la cartera, identificar atrasos, calcular cuotas, organizar visitas y demostrar pagos. En campo también se necesita trabajar con conectividad limitada y utilizar GPS, cámara, QR, BLE y NFC.")
d.add_heading("1.2 Objetivo general",2)
d.add_paragraph("Desarrollar una aplicación móvil híbrida con Ionic, Angular y Capacitor para administrar clientes, préstamos, cuotas, pagos, tareas y comunicaciones, integrando almacenamiento local, servicios web y capacidades nativas.")
d.add_heading("1.3 Objetivos específicos",2)
bullets(["Implementar navegación protegida con pestañas y múltiples pantallas.","Gestionar clientes, préstamos, pagos y tareas con CRUD persistente.","Calcular cuotas, vencimientos, saldos y estados automáticamente.","Integrar conectividad, GPS y mapa interactivo de clientes.","Incorporar audio, cámara, QR, BLE y NFC con fallbacks web.","Consumir servicios web con loading, timeout, errores y respaldo.","Generar y compartir recibos por PDF, WhatsApp, correo y NFC."])
d.add_heading("1.4 Alcance",2)
table(["Incluye","No incluye en 1.0"],[("Sesión, perfil y empresa.","Servidor multiempresa y recuperación remota."),("Clientes, préstamos, cuotas, pagos e historial.","Buró de crédito y conciliación bancaria."),("Tareas, recordatorios, correo y WhatsApp.","Envío silencioso de WhatsApp."),("GPS, mapa y geocodificación.","Navegación giro a giro."),("Persistencia local y cola offline.","Sincronización productiva con backend.")])
d.add_heading("1.5 Público objetivo",2);d.add_paragraph("Prestamistas independientes, cobradores, microempresas de crédito y administradores de carteras pequeñas en República Dominicana.")

d.add_heading("2. Matriz de cumplimiento del curso",1)
table(["Unidad","Tema","Evidencia","Estado"],[("1–2","Introducción/navegación","IonTabs, Router, AuthGuard, 13 pantallas y carga organizada.","Cumplido"),("3","Interfaces/gestos","Cards, listas, inputs, modales, swipe, refresher, long press y reorder.","Cumplido"),("4","Conectividad","Network, online/offline, tipo y cola local.","Cumplido"),("5","Bluetooth/NFC","Escaneo/conexión BLE y lectura/escritura NFC.","Cumplido"),("6","Geolocalización","GPS, Leaflet, marcadores, geocodificación y distancia.","Cumplido"),("7","Multimedia","Audio con controles, progreso y volumen.","Cumplido"),("8","Cámara/QR","Camera, ML Kit, BarcodeDetector y getUserMedia.","Cumplido"),("9","Almacenamiento","Ionic Storage/localStorage y CRUD.","Cumplido"),("10","Servicios web","HttpClient, fetch, EmailJS, errores y fallback.","Cumplido")],7.8)
capture("R-01","Dashboard mostrando accesos y barra de tabs.")

d.add_heading("3. Modelo de negocio — Canvas",1)
table(["Bloque","Aplicación a FlexCrediAPP"],[("Segmentos","Prestamistas, cobradores, microfinancieras y pequeños negocios."),("Propuesta de valor","Control móvil de cartera, menos errores, recibos inmediatos, mapa y comunicación."),("Canales","APK/Google Play, futura App Store y web."),("Relación","Autoservicio, soporte, actualizaciones y feedback."),("Ingresos","Freemium o suscripción con nube y multiusuario."),("Recursos","Código, plugins, APIs, dispositivos y conocimiento financiero."),("Actividades","Desarrollo, QA multidispositivo, mantenimiento y soporte."),("Socios","Ionic/Capacitor, Google/Apple, OSM, EmailJS y WhatsApp."),("Costos","Desarrollo, pruebas, publicación, nube, soporte y mantenimiento.")])
capture("C-01","Convierta los nueve bloques en un lienzo Canvas visual de una página.","DIAGRAMA")

d.add_heading("4. Arquitectura técnica",1)
d.add_paragraph("La aplicación usa componentes Angular standalone. Las páginas coordinan casos de uso; los servicios encapsulan negocio, persistencia e integraciones; los modelos definen contratos; Capacitor enlaza web con Android/iOS. Cada capacidad nativa contempla soporte o mensaje alternativo en navegador.")
table(["Capa","Responsabilidad","Elementos"],[("Presentación","UI, navegación y feedback.","pages, tabs, HTML/SCSS, Ionic, Ionicons."),("Aplicación","Coordinar registros, cobros y compartidos.","Clases de página y AuthGuard."),("Dominio","Entidades y reglas financieras.","Models y DataService."),("Servicios","Persistencia, red, GPS, cámara, QR, BLE, NFC, correo.","15 servicios Angular."),("Integración","Plugins y APIs externas.","Capacitor, JSONPlaceholder, Nominatim, EmailJS, wa.me."),("Datos","Persistencia y cola offline.","localStorage e Ionic Storage.")],8)
d.add_heading("4.1 Diagrama lógico",2)
code("USUARIO\n  ↓\nPANTALLAS IONIC + ION TABS + ROUTER\n  ↓\nSERVICIOS DE NEGOCIO ── SERVICIOS NATIVOS\n  ↓                         ↓\nLOCAL STORAGE / IONIC STORAGE     CAPACITOR\n  ↓                         ↓\nAPIs REST / EmailJS / WhatsApp    Android / iOS", "Arquitectura resumida")
capture("A-01","Diagrama gráfico Vista → Página → Servicio → Storage/API/Plugin.","DIAGRAMA")
d.add_heading("4.2 Stack",2)
table(["Tecnología","Versión","Uso"],[("Ionic Angular","8.x","UI móvil y navegación."),("Angular","20.x","Componentes, formularios, routing y HTTP."),("Capacitor","8.4.x","Puente Android/iOS."),("TypeScript","5.9.x","Lenguaje tipado."),("Leaflet","1.9.4","Mapa."),("RxJS","7.8.x","Observables y HTTP."),("Ionic Storage","4.x","Persistencia asíncrona."),("Jasmine/Karma/ESLint","5.1/6.4/9.16","Pruebas y calidad.")])
d.add_heading("4.3 Plugins",2)
table(["Plugin","Versión","Uso"],[("@capacitor-community/bluetooth-le","8.2.0","Escaneo/conexión BLE y RSSI."),("@capacitor-mlkit/barcode-scanning","8.1.0","QR nativo."),("@capacitor/camera","8.2.1","Cámara/galería."),("@capacitor/geolocation","8.2.0","GPS/permisos."),("@capacitor/network","8.0.1","Red y listeners."),("@capgo/capacitor-nfc","8.2.2","Lectura/escritura NFC."),("@capacitor/app","8.1.0","Ciclo de vida."),("@capacitor/haptics","8.0.2","Respuesta háptica."),("@capacitor/keyboard","8.0.3","Teclado móvil."),("@capacitor/status-bar","8.0.2","Barra de estado.")],7.8)
d.add_heading("4.4 Carpetas y flujo",2)
code("src/app/\n├── guards/ y models/\n├── pages/ (13 pantallas)\n├── services/ (15 servicios)\n└── tabs/\nandroid/  ios/  src/assets/", "Estructura principal")
capture("A-02","Árbol de VS Code mostrando pages, services, models, tabs, android e ios.","CÓDIGO")
numbers(["RecibirPagoPage obtiene préstamos activos.","El usuario selecciona cliente y cuota.","DataService guarda pago y actualiza cuota/préstamo.","Se muestra recibo y saldo.","Email, WhatsApp, PDF y NFC comparten evidencia.","Dashboard, historial y reportes leen el mismo estado."])

d.add_heading("5. Diseño de interfaces",1)
bullets(["Jerarquía visual con azul oscuro, acciones azules/verdes y alertas rojas.","Objetivos táctiles amplios, cards, modales, estados vacíos y confirmaciones.","Iconografía Ionicons, bordes redondeados, sombras y tipografía legible.","Spinners, toasts, badges, estados de red y validaciones.","Etiquetas, botones semánticos, contraste y texto alternativo."])
table(["Uso","Color","Código"],[("Identidad","Azul oscuro","#0D2252"),("Acción","Azul real","#1A56DB"),("Acento","Dorado","#F59E0B"),("Éxito","Verde","#059669"),("Peligro","Rojo","#DC2626"),("Fondos","Lavanda/gris","#EEF2FF / #F8FAFC")])
d.add_heading("5.1 Navegación",2)
code("Login → Tabs\nTabs → Inicio | Tareas | Novedades | Mapa | Perfil\nInicio → Clientes → Nuevo/Editar cliente\nInicio → Nuevo préstamo\nInicio → Recibir pago → QR/Recibo/PDF/WhatsApp/NFC\nInicio → Historial | Reportes | Podcasts", "Mapa de navegación")
capture("D-01","Mapa de navegación convertido en diagrama profesional.","DIAGRAMA")
capture("D-02","Wireframes o mosaico de Login, Dashboard, Tareas, Mapa, Recibir pago y Perfil.")

d.add_heading("6. Desarrollo por módulos",1)
d.add_paragraph("Sustituya Estudiante A–E por nombres y matrículas reales. Cada módulo documenta necesidad, solución, decisiones, dificultades, código y capturas.")

def module(num,title,responsable,unidades,problema,solucion,componentes,decisiones,dificultades,codigo,codigo_titulo,capturas):
    d.add_page_break(); d.add_heading(f"6.{num} Módulo {num}: {title}",1)
    table(["Responsable","Unidades"],[(responsable,unidades)],9)
    d.add_heading("Problema a resolver",2);d.add_paragraph(problema)
    d.add_heading("Solución implementada",2);d.add_paragraph(solucion)
    d.add_heading("Componentes, servicios y funciones",2);table(["Elemento","Implementación"],componentes,8)
    d.add_heading("Decisiones técnicas",2);bullets(decisiones)
    code(codigo,codigo_titulo)
    d.add_heading("Dificultades y soluciones",2);d.add_paragraph(dificultades)
    for c in capturas:capture(*c)

module(1,"Navegación, autenticación y novedades","[ESTUDIANTE A — NOMBRE Y MATRÍCULA]","U1, U2 y U10",
"El usuario necesita entrar, desplazarse con un solo toque y consultar contenido aunque falle Internet.",
"Se implementaron registro/login local, AuthGuard, IonTabs, rutas standalone y Novedades con API REST, spinner, timeout, pull-to-refresh, reintento y respaldo local.",
[("Navegación","IonTabs, IonTabBar, IonTabButton, Router, redirects y AuthGuard."),("API","GET JSONPlaceholder /posts y /posts/{id}."),("Resiliencia","Timeout 8 s, catchError, noticias locales y finalización del refresher."),("Dashboard","Cartera, clientes, atrasos, recordatorios y accesos."),("Páginas","Inicio, Tareas, Novedades, Mapa, Podcasts, Perfil y flujos financieros.")],
["HttpClient/RxJS separan HTTP de la vista y facilitan map, timeout y fallback.","El respaldo local conserva valor sin conexión.","Las rutas se organizan por tabs y pantallas protegidas."],
"return this.http.get<any[]>(this.API_URL).pipe(\n  timeout(8000), delay(600),\n  map(posts => posts.slice(0,20).map(p => this.transformar(p))),\n  catchError(() => of(this.noticiasLocales()))\n);","Código relevante: consumo REST",
"Se corrigieron loaders persistentes, codificación de textos y una transición hacia Reportes que retenía gestos. Se garantizó cierre en éxito/error y se ajustó la navegación Ionic.",
[("M1-01","Login con registro, acceso y validaciones."),("M1-02","Dashboard completo y tabs."),("M1-03","Novedades y pull-to-refresh."),("M1-C1","app.routes.ts: AuthGuard, tabs, rutas hijas y lazy loading.","CÓDIGO"),("M1-C2","NoticiasService: HttpClient, timeout, map y catchError.","CÓDIGO")])

module(2,"Tareas, interfaces, gestos y almacenamiento","[ESTUDIANTE B — NOMBRE Y MATRÍCULA]","U3 y U9",
"El cobrador necesita organizar llamadas, visitas y cobros por prioridad, cliente y fecha.",
"Tareas permite crear, leer, editar, completar, cancelar, reabrir, eliminar, filtrar y reordenar. Persiste con Ionic Storage.",
[("CRUD","TareasService: agregar, obtener, actualizar y eliminar."),("Swipe","ion-item-sliding con editar/cancelar/eliminar."),("Pull-to-refresh","ion-refresher recarga colección."),("Long press","Temporizador abre edición."),("Reorder","ion-reorder-group persiste orden."),("Datos","Título, descripción, cliente, prioridad, fecha, tipo y estado.")],
["Ionic Storage ofrece API asíncrona y desacopla la persistencia.","stopPropagation evita conflictos entre gestos y botones.","El orden se conserva explícitamente al reorganizar."],
"async agregarTarea(datos) {\n  const lista = await this.leer();\n  const nueva = { ...datos, id: `TAR-${Date.now()}`, orden: lista.length };\n  lista.push(nueva); await this.guardar(lista); return nueva;\n}","Código relevante: Create del CRUD",
"El reto fue distinguir toque y pulsación prolongada, evitar eventos duplicados y mantener reorder al filtrar. Se usó temporizador cancelable y control de propagación.",
[("M2-01","Tareas con estados y estadísticas."),("M2-02","Swipe mostrando opciones."),("M2-03","Modal crear/editar completo."),("M2-C1","HTML de ion-item-sliding e ion-reorder.","CÓDIGO"),("M2-C2","TareasService con CRUD y persistencia.","CÓDIGO")])

module(3,"Conectividad, GPS y mapa de clientes","[ESTUDIANTE C — NOMBRE Y MATRÍCULA]","U4 y U6",
"La cobranza de campo requiere saber si existe red, ubicar al usuario y visualizar clientes por zona.",
"NetworkService publica conexión/tipo; MapaPage usa Leaflet/OSM, GPS, Nominatim, marcadores personalizados, popups y distancia Haversine.",
[("Red","@capacitor/network, BehaviorSubject y fallback online/offline."),("GPS","Permiso, alta precisión, timeout 20 s y navigator.geolocation."),("Mapa","Leaflet + mosaicos OpenStreetMap."),("Clientes","Marcadores con iniciales, contacto, dirección y préstamos."),("Geocodificación","Nominatim countrycodes=do."),("Proximidad","Clientes dentro de 3 km por Haversine.")],
["BehaviorSubject actualiza cualquier vista suscrita.","OpenStreetMap evita una clave comercial para la demostración.","Fallback web permite probar sin dispositivo nativo."],
"const pos = await Geolocation.getCurrentPosition({\n  enableHighAccuracy: true, timeout: 20000, maximumAge: 0\n});\nreturn { lat: pos.coords.latitude, lng: pos.coords.longitude };","Código relevante: GPS",
"La ubicación varió entre navegadores y dispositivos por HTTPS, permisos y precisión. En algunos falló y en otros funcionó excelentemente con marcadores correctos. Se añadieron timeout, fallback, mensajes, centro RD y geocodificación progresiva.",
[("M3-01","Indicador online/offline y modal de red."),("M3-02","Mapa con ubicación y marcadores."),("M3-03","Popup de cliente."),("M3-C1","NetworkService y listener de cambios.","CÓDIGO"),("M3-C2","Leaflet tileLayer y marcadores.","CÓDIGO")])

module(4,"Multimedia, cámara y QR","[ESTUDIANTE D — NOMBRE Y MATRÍCULA]","U7 y U8",
"El proyecto requiere contenido multimedia y captura/escaneo desde el dispositivo aplicados al contexto financiero.",
"Podcasts usa HTML5 Audio; CameraService gestiona cámara/galería; QrService selecciona ML Kit nativo o BarcodeDetector web para localizar clientes en Recibir pago.",
[("Audio","Play, pausa, stop, anterior, siguiente, progreso, seek y volumen."),("Cámara","@capacitor/camera, DataUrl, calidad 80 y ancho 1024."),("Web","getUserMedia o input file."),("QR nativo","ML Kit Barcode Scanning."),("QR web","BarcodeDetector, video y cierre de tracks."),("Aplicación","Autoselección de préstamo por QR.")],
["El audio se inicia por gesto para respetar autoplay.","La cámara retorna Data URL para persistencia local.","Se separan rutas nativas/web y se detiene MediaStream al cerrar."],
"audio.addEventListener('timeupdate', () => {\n  this.currentTime = audio.currentTime;\n  this.duration = audio.duration || 0;\n  this.progreso = this.duration ? this.currentTime / this.duration * 100 : 0;\n});","Código relevante: progreso de audio",
"Cámara y QR dependen de permisos, iluminación, contexto seguro y soporte. Se incorporaron fallback, cancelación segura, mensajes y pruebas en Android físico.",
[("M4-01","Catálogo de podcasts."),("M4-02","Reproductor y controles."),("M4-03","Escáner QR en Recibir pago."),("M4-04","Foto/logo desde cámara o galería."),("M4-C1","Eventos de Audio API.","CÓDIGO"),("M4-C2","QrService nativo/web.","CÓDIGO")])

module(5,"Bluetooth, NFC, perfil y configuración","[ESTUDIANTE E — NOMBRE Y MATRÍCULA]","U5 y U9",
"La aplicación debe aprovechar tecnologías de proximidad y conservar identidad/configuración del prestamista.",
"Perfil gestiona datos, foto, logo, tema y toggles; BluetoothService escanea/conecta BLE y muestra RSSI; NfcService lee/escribe texto o recibos y registra historial.",
[("Perfil","Nombre, empresa, correo, teléfono, documento, profesión y dirección."),("Configuración","Tema, notificaciones, sonidos, sincronización y vibración."),("BLE","Plugin nativo y Web Bluetooth cuando aplica."),("NFC","Lectura/escritura, estado y apertura de ajustes."),("Historial","Tipo, destino, contenido, fecha y éxito."),("Persistencia","Ionic Storage para configuración, fotos y compartidos.")],
["Se detecta plataforma antes de invocar hardware.","RSSI se traduce en una señal comprensible.","La interfaz informa cuando una capacidad no está soportada."],
"const dispositivos = await this.bluetooth.escanear();\n// Resultado: id, nombre y RSSI.\n// La UI permite seleccionar y conectar.","Código relevante: flujo BLE",
"BLE/NFC requieren hardware, permisos y dispositivo real. Se implementó detección honesta y se reservaron pruebas finales para Android físico compatible.",
[("M5-01","Perfil y configuraciones."),("M5-02","Modal BLE con RSSI y Conectar."),("M5-03","NFC e historial."),("M5-C1","BluetoothService nativo/web.","CÓDIGO"),("M5-C2","NfcService y listeners.","CÓDIGO")])

d.add_page_break();d.add_heading("7. Funcionalidades adicionales",1)
d.add_paragraph("Estas funciones exceden el ejemplo mínimo y forman el flujo real del producto.")
table(["Función","Detalle"],[("Autenticación","Registro, login, logout, AuthGuard y perfil."),("Clientes","CRUD, búsqueda, historial, coordenadas y totales."),("Préstamos","Cliente, monto, interés, frecuencia, cuotas, fechas y estados."),("Cobros","Selección/QR, confirmación, pago de cuota y cierre del préstamo."),("Recibo","Número, logo, cliente, cuota, saldo, fecha y estado de correo."),("PDF","Generación en memoria, descarga y Web Share."),("WhatsApp","Normalización dominicana, recordatorios y recibo por wa.me."),("EmailJS","Plantilla parametrizada y confirmación automática."),("Historial","Vista global/cliente, totales y relación con préstamo."),("Reportes","Semana/mes/año, KPIs y gráfica semanal."),("Recordatorios","Pendientes por fecha y control de enviados.")],7.8)
d.add_heading("7.1 Servicios externos",2)
table(["Servicio","Operación","Uso","Fallos"],[("JSONPlaceholder","GET /posts","Noticias financieras adaptadas.","Timeout y fallback local."),("Nominatim","GET /search","Dirección a coordenadas.","Variantes y omisión segura."),("OpenStreetMap","Tiles HTTPS","Mapa base.","Dependencia de red."),("EmailJS 4","send","Correo de pago.","Validación y try/catch."),("WhatsApp","wa.me","Mensaje preparado.","Normalización/fallback."),("SoundHelix/Picsum","GET","Audio/portadas demo.","Spinner y toast.")],7.5)
note("SEGURIDAD","No incluya claves, contraseñas, cédulas ni teléfonos reales en capturas públicas. En producción, use variables de entorno y backend.",True)
for c in [("EX-01","Cliente, cuota y Registrar pago."),("EX-02","Recibo con PDF/WhatsApp/NFC."),("EX-03","Correo recibido o evidencia EmailJS sin datos sensibles."),("EX-04","WhatsApp con mensaje prellenado y datos ocultos."),("EX-05","Historial y Reportes con KPIs.")]:capture(*c)

d.add_heading("8. Persistencia y modelo de datos",1)
table(["Entidad","Campos","Relación"],[("Usuario","id, nombre, email, empresa, foto y datos.","Sesión y configuración."),("Cliente","id, identidad, contacto, dirección y coordenadas.","1:N préstamos."),("Préstamo","clienteId, monto, interés, frecuencia, cuotas, estado.","1:N cuotas/pagos."),("Cuota","número, vencimiento, monto, estado, fechaPago.","Pertenece a préstamo."),("Pago","préstamo, cliente, cuota, monto y fecha.","Actualiza cuota/reportes."),("Tarea","título, cliente, prioridad, fecha, estado, orden.","0..N por cliente."),("Compartido","tipo, destino, contenido, fecha, éxito.","Auditoría BLE/NFC/QR.")],8)
table(["Clave","Contenido","Mecanismo"],[("fc_clientes","Clientes","localStorage"),("fc_prestamos","Préstamos/cuotas","localStorage"),("fc_pagos","Pagos","localStorage"),("fc_sync_queue","Cola offline","localStorage"),("flexcredi_tareas","Tareas","Ionic Storage"),("config/foto/logo/compartidos","Perfil","Ionic Storage")])
capture("DB-01","Diagrama ER: Cliente 1:N Préstamo; Préstamo 1:N Cuota; Pago relacionado con cliente/préstamo/cuota; Tarea opcionalmente asociada.","DIAGRAMA")

d.add_heading("9. Pruebas y despliegue",1)
d.add_paragraph("Se ejecutaron revisión estática, compilación, prueba unitaria, sincronización Capacitor y pruebas manuales. Hardware debe validarse en teléfono real.")
tests=[("Registro/login","Sesión y AuthGuard.","Pasó"),("Tabs/rutas","Un toque y retorno.","Pasó"),("CRUD cliente","Persistencia y conteos.","Pasó"),("Crear préstamo","Cuotas/calendario.","Pasó"),("Registrar pago","Cuota, recibo, historial.","Pasó"),("QR","Selecciona coincidencia.","Android"),("PDF","Genera/descarga/comparte.","Pasó"),("WhatsApp","Chat y mensaje.","Pasó"),("EmailJS","Correo y estado.","Con red"),("Tareas CRUD","Persistencia.","Pasó"),("Gestos","Swipe/long/reorder/refresher.","Pasó"),("Noticias online/offline","API/fallback.","Pasó"),("Red","Online/offline/tipo.","Pasó"),("Mapa","OSM/marcadores/popups.","Con red"),("GPS","Permiso/precisión.","Variable"),("Podcasts","Controles completos.","Con red"),("Cámara","Foto/galería.","Android"),("BLE","Escaneo/conexión.","Android físico"),("NFC","Lectura/escritura.","Android físico"),("Perfil","Datos/tema/foto.","Pasó"),("Lint/build/test","Sin errores.","Pasó")]
table(["Prueba","Evidencia esperada","Estado"],tests,7.8)
d.add_heading("9.1 Bugs corregidos",2)
table(["Hallazgo","Causa","Solución"],[("Loader permanecía","Ciclo incompleto.","Cerrar en éxito/error."),("Recibir pago vacío","Bloque visual retirado.","Restaurar lista, cuota, footer y QR."),("Reportes con doble toque","Transición tabs/página.","Navegación Ionic sin animación."),("Textos dañados","Codificación mojibake.","Normalización UTF-8."),("GPS inconsistente","HTTPS/permisos/soporte.","Fallback, timeout y mensajes.")],7.8)
for c in [("T-01","npm run lint exitoso.","PRUEBA"),("T-02","npm run build exitoso.","PRUEBA"),("T-03","Karma/Jasmine TOTAL SUCCESS.","PRUEBA"),("T-04","App en ionic serve.","PRUEBA"),("T-05","App en emulador Android Studio.","PRUEBA"),("T-06","App instalada en Android real.","PRUEBA")]:capture(*c)
d.add_heading("9.2 Instalación",2)
numbers(["Instalar Node.js, Git, Ionic/Angular CLI y Android Studio.","Abrir el proyecto y ejecutar npm install.","Navegador: npm start o ionic serve.","Nativos: npm run build y npx cap sync.","Android: npx cap open android.","Generar APK debug con Build > Build APK(s) o gradlew assembleDebug."])
d.add_heading("9.3 Condiciones",2)
bullets(["GPS: permiso, ubicación activa y HTTPS en navegador.","Cámara/QR: permiso e iluminación.","BLE: Bluetooth y permisos cercanos.","NFC: hardware y NFC activo.","Noticias, mapa, correo y audio: Internet."])

d.add_heading("10. Seguridad, privacidad y limitaciones",1)
table(["Tema","Actual","Producción"],[("Credenciales","Login académico local.","Backend, hash y tokens."),("Finanzas","Datos en dispositivo.","Cifrado, copias y roles."),("APIs","Cliente llama servicios.","Variables/proxy/backend."),("Privacidad","Identidad, ubicación y pagos.","Consentimiento y retención."),("Offline","Cola preparada/simulada.","Sync idempotente y conflictos."),("Hardware","Depende de compatibilidad.","Matriz oficial de equipos.")],8)

d.add_heading("11. Dificultades, lecciones y conclusiones",1)
d.add_heading("11.1 Dificultades",2)
bullets(["Pruebas en navegadores/dispositivos con soporte y permisos diferentes.","GPS falló en algunos entornos y funcionó excelentemente en otros.","BLE/NFC/QR/cámara requieren pruebas físicas reales.","Tabs y rutas independientes exigieron depuración táctil.","Un pago debe actualizar cuota, préstamo, recibo, historial y reporte.","Se revisaron acentos, loaders, barras y alineación profesional."])
d.add_heading("11.2 Reflexión",2)
d.add_paragraph("Crear una app híbrida es un reto grande: una base de código llega a varias plataformas, pero cada entorno cambia permisos, sensores, ciclo de vida, seguridad y rendimiento. Por ello estas aplicaciones requieren tiempo de desarrollo y muchas pruebas en navegadores, emuladores y dispositivos físicos. La calidad exige probar permisos denegados, pérdida de red, tamaños de pantalla y fallbacks.")
d.add_heading("11.3 Conclusión",2)
d.add_paragraph("FlexCrediAPP integra las diez unidades en un producto coherente. La navegación conecta procesos financieros reales; los servicios móviles mejoran la cobranza; el almacenamiento brinda continuidad; y las comunicaciones convierten un pago en evidencia. El proyecto fortaleció arquitectura, Angular standalone, Ionic, persistencia, APIs, sensores, errores y QA multiplataforma.")
d.add_heading("11.4 Versión 2.0",2)
bullets(["Backend multiempresa, roles, auditoría y copias.","Base cifrada y autenticación segura.","Push y recordatorios en segundo plano.","Reportes exportables y riesgo.","Pruebas E2E, accesibilidad, telemetría y CI/CD.","Publicación firmada y matriz de dispositivos."])

d.add_heading("12. Referencias APA 7",1)
refs=["Angular. (2026). Angular documentation. https://angular.dev/","Capacitor. (2026). Capacitor documentation. https://capacitorjs.com/docs","EmailJS. (2026). EmailJS documentation. https://www.emailjs.com/docs/","Google. (2026). ML Kit barcode scanning. https://developers.google.com/ml-kit/vision/barcode-scanning","Ionic. (2026). Ionic Framework documentation. https://ionicframework.com/docs","Leaflet. (2025). Leaflet. https://leafletjs.com/","MDN Web Docs. (2025). HTMLMediaElement. https://developer.mozilla.org/docs/Web/API/HTMLMediaElement","MDN Web Docs. (2025). MediaDevices: getUserMedia(). https://developer.mozilla.org/docs/Web/API/MediaDevices/getUserMedia","OpenStreetMap contributors. (2026). OpenStreetMap. https://www.openstreetmap.org/","OpenStreetMap Foundation. (2026). Nominatim documentation. https://nominatim.org/release-docs/latest/","ReactiveX. (2026). RxJS documentation. https://rxjs.dev/","Strategyzer. (2025). Business Model Canvas. https://www.strategyzer.com/library/the-business-model-canvas","TypeScript. (2026). TypeScript documentation. https://www.typescriptlang.org/docs/","WhatsApp. (2026). Click to chat. https://faq.whatsapp.com/5913398998672934/","World Wide Web Consortium. (2023). WCAG 2.2. https://www.w3.org/TR/WCAG22/"]
for ref in refs:
    p=d.add_paragraph(ref);p.paragraph_format.left_indent=Inches(.3);p.paragraph_format.first_line_indent=Inches(-.3)

d.add_heading("Anexos",1)
d.add_heading("Anexo A. Manual de usuario",2)
numbers(["Crear cuenta e iniciar sesión.","Registrar cliente.","Crear préstamo y revisar calendario.","Consultar Dashboard.","Registrar pago por lista o QR.","Compartir recibo por PDF, WhatsApp, correo o NFC.","Consultar historial/reportes.","Gestionar tareas y gestos.","Usar mapa, podcasts, perfil, BLE y NFC."])
d.add_heading("Anexo B. Archivos para la defensa",2)
table(["Archivo","Qué explicar"],[("app.routes.ts","Tabs, guard, rutas y lazy loading."),("data.service.ts","CRUD, cuotas, pagos y estadísticas."),("tareas.service.ts","Storage y CRUD."),("network.service.ts","Observable y cola offline."),("geolocation.service.ts / mapa.page.ts","GPS, geocodificación y Leaflet."),("qr.service.ts / camera.service.ts","Nativo vs web."),("bluetooth.service.ts / nfc.service.ts","Hardware y permisos."),("recibir-pago.page.ts","Pago, recibo, correo, PDF, WhatsApp y NFC.")],8)
d.add_heading("Anexo C. Lista maestra de capturas",2)
table(["Sección","Códigos","Contenido"],[("Portada","P-01","Logo UAPA"),("Arquitectura","A-01/A-02","Diagrama/árbol"),("Diseño","D-01/D-02","Navegación/wireframes"),("Módulo 1","M1-*","Login/dashboard/API"),("Módulo 2","M2-*","Tareas/gestos/storage"),("Módulo 3","M3-*","Red/GPS/mapa"),("Módulo 4","M4-*","Audio/cámara/QR"),("Módulo 5","M5-*","Perfil/BLE/NFC"),("Extras","EX-*","Pago/recibo/correo/WA/reportes"),("Pruebas","T-*","Lint/build/test/dispositivos"),("Datos","DB-01","Modelo ER")],8)
d.add_heading("Anexo D. Guion de demostración",2)
numbers(["Login y tabs.","Cliente y préstamo.","Tarea con gestos.","Red y Novedades.","GPS/mapa.","Podcast.","Cámara/QR.","BLE/NFC.","Pago, recibo, historial y reportes."])
note("REVISIÓN FINAL","Complete datos, inserte capturas, actualice índice, numere figuras/tablas, revise ortografía y oculte datos sensibles.")

d.core_properties.title="Documentación técnica y funcional de FlexCrediAPP"
d.core_properties.subject="Proyecto final — ISW-307"
d.core_properties.author="Equipo FlexCrediAPP"
d.save(OUT)
print(OUT)
